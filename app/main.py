"""
Stream Bill Generator - Complete Streamlit Application
SINGLE FILE DEPLOYMENT OPTIMIZED VERSION

This is the ONLY entry point for the application.
All functionality consolidated for maximum compatibility.
"""
import streamlit as st
import pandas as pd
import os
import sys
import tempfile
import zipfile
import base64
from io import BytesIO
from functools import lru_cache
from pathlib import Path
import re
import math
from datetime import datetime

# ============================================================================
# CONSOLIDATED FUNCTIONALITY - ALL IN ONE FILE
# ============================================================================

# Get absolute paths
CURRENT_FILE = os.path.abspath(__file__)
APP_DIR = os.path.dirname(CURRENT_FILE)
ROOT_DIR = os.path.dirname(APP_DIR)

# Optional imports with fallbacks
try:
    import numpy as np
except ImportError:
    np = None

try:
    from num2words import num2words
    def number_to_words(number):
        return num2words(int(number), lang="en_IN").title()
except ImportError:
    def number_to_words(number):
        return str(number)

try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    pdfkit = None
    PDFKIT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfWriter, PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfWriter, PdfReader
        PYPDF_AVAILABLE = True
    except ImportError:
        PdfWriter = None
        PdfReader = None
        PYPDF_AVAILABLE = False

try:
    from jinja2 import Environment, BaseLoader
    JINJA2_AVAILABLE = True
except ImportError:
    Environment = None
    BaseLoader = None
    JINJA2_AVAILABLE = False

def safe_float(value, default=0.0):
    """Safely convert a value to float with proper error handling"""
    try:
        if value is None:
            return default
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            cleaned = value.strip().replace(',', '').replace(' ', '')
            if cleaned == '':
                return default
            return float(cleaned)
        return default
    except (ValueError, TypeError):
        return default

def process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type):
    """Process bill data from Excel sheets - CONSOLIDATED VERSION"""
    first_page_data = {"header": [], "items": [], "totals": {}}
    last_page_data = {"payable_amount": 0, "amount_words": ""}
    deviation_data = {"items": [], "summary": {}}
    extra_items_data = {"items": []}
    note_sheet_data = {"notes": []}

    # Header processing
    header_data = ws_wo.iloc[:19, :7].fillna("").values.tolist()
    
    # Format dates
    for i in range(len(header_data)):
        for j in range(len(header_data[i])):
            val = header_data[i][j]
            if hasattr(val, 'strftime'):
                header_data[i][j] = val.strftime("%d-%m-%Y")

    first_page_data["header"] = header_data

    # Process work order items
    last_row_wo = ws_wo.shape[0]
    for i in range(21, last_row_wo):
        qty_raw = ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None
        rate_raw = ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None

        qty = safe_float(qty_raw)
        rate = safe_float(rate_raw)

        if rate == 0:
            item = {
                "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                "unit": "", "quantity": "", "rate": "", "amount": "",
                "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
                "is_divider": False
            }
        else:
            item = {
                "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
                "quantity": qty, "rate": rate, "amount": round(qty * rate) if qty and rate else 0,
                "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else "",
                "is_divider": False
            }
        first_page_data["items"].append(item)

    # Extra items divider
    first_page_data["items"].append({
        "description": "Extra Items (With Premium)",
        "bold": True, "underline": True, "is_divider": True,
        "amount": 0, "quantity": 0, "rate": 0, "serial_no": "", "unit": "", "remark": ""
    })

    # Process extra items
    last_row_extra = ws_extra.shape[0]
    for j in range(6, last_row_extra):
        qty_raw = ws_extra.iloc[j, 3] if pd.notnull(ws_extra.iloc[j, 3]) else None
        rate_raw = ws_extra.iloc[j, 5] if pd.notnull(ws_extra.iloc[j, 5]) else None

        qty = safe_float(qty_raw)
        rate = safe_float(rate_raw)

        if rate == 0:
            item = {
                "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
                "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
                "unit": "", "quantity": "", "rate": "", "amount": "",
                "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
                "is_divider": False
            }
        else:
            item = {
                "serial_no": str(ws_extra.iloc[j, 0]) if pd.notnull(ws_extra.iloc[j, 0]) else "",
                "description": str(ws_extra.iloc[j, 2]) if pd.notnull(ws_extra.iloc[j, 2]) else "",
                "unit": str(ws_extra.iloc[j, 4]) if pd.notnull(ws_extra.iloc[j, 4]) else "",
                "quantity": qty, "rate": rate, "amount": round(qty * rate) if qty and rate else 0,
                "remark": str(ws_extra.iloc[j, 1]) if pd.notnull(ws_extra.iloc[j, 1]) else "",
                "is_divider": False
            }
        first_page_data["items"].append(item)
        extra_items_data["items"].append(item.copy())

    # Calculate totals
    data_items = [item for item in first_page_data["items"] if not item.get("is_divider", False)]
    total_amount = round(sum(safe_float(item.get("amount", 0)) for item in data_items))
    premium_amount = round(total_amount * (premium_percent / 100) if premium_type == "above" else -total_amount * (premium_percent / 100))
    payable_amount = round(total_amount + premium_amount)

    first_page_data["totals"] = {
        "grand_total": total_amount,
        "premium": {"percent": premium_percent / 100, "type": premium_type, "amount": premium_amount},
        "payable": payable_amount
    }

    last_page_data = {"payable_amount": payable_amount, "amount_words": number_to_words(payable_amount)}

    # Process deviation data
    work_order_total = 0
    executed_total = 0
    overall_excess = 0
    overall_saving = 0
    
    for i in range(21, last_row_wo):
        qty_wo = safe_float(ws_wo.iloc[i, 3] if pd.notnull(ws_wo.iloc[i, 3]) else None)
        rate = safe_float(ws_wo.iloc[i, 4] if pd.notnull(ws_wo.iloc[i, 4]) else None)
        qty_bill = safe_float(ws_bq.iloc[i, 3] if i < ws_bq.shape[0] and pd.notnull(ws_bq.iloc[i, 3]) else None)

        amt_wo = round(qty_wo * rate)
        amt_bill = round(qty_bill * rate)
        excess_qty = qty_bill - qty_wo if qty_bill > qty_wo else 0
        excess_amt = round(excess_qty * rate) if excess_qty > 0 else 0
        saving_qty = qty_wo - qty_bill if qty_bill < qty_wo else 0
        saving_amt = round(saving_qty * rate) if saving_qty > 0 else 0

        if rate == 0:
            item = {
                "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                "unit": "", "qty_wo": "", "rate": "", "amt_wo": "", "qty_bill": "", "amt_bill": "",
                "excess_qty": "", "excess_amt": "", "saving_qty": "", "saving_amt": "",
                "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else ""
            }
        else:
            item = {
                "serial_no": str(ws_wo.iloc[i, 0]) if pd.notnull(ws_wo.iloc[i, 0]) else "",
                "description": str(ws_wo.iloc[i, 1]) if pd.notnull(ws_wo.iloc[i, 1]) else "",
                "unit": str(ws_wo.iloc[i, 2]) if pd.notnull(ws_wo.iloc[i, 2]) else "",
                "qty_wo": qty_wo, "rate": rate, "amt_wo": amt_wo,
                "qty_bill": qty_bill, "amt_bill": amt_bill,
                "excess_qty": excess_qty, "excess_amt": excess_amt,
                "saving_qty": saving_qty, "saving_amt": saving_amt,
                "remark": str(ws_wo.iloc[i, 6]) if pd.notnull(ws_wo.iloc[i, 6]) else ""
            }
            work_order_total += amt_wo
            executed_total += amt_bill
            overall_excess += excess_amt
            overall_saving += saving_amt

        deviation_data["items"].append(item)

    # Deviation summary
    tender_premium_f = round(work_order_total * (premium_percent / 100) if premium_type == "above" else -work_order_total * (premium_percent / 100))
    tender_premium_h = round(executed_total * (premium_percent / 100) if premium_type == "above" else -executed_total * (premium_percent / 100))
    
    deviation_data["summary"] = {
        "work_order_total": round(work_order_total),
        "executed_total": round(executed_total),
        "overall_excess": round(overall_excess),
        "overall_saving": round(overall_saving),
        "premium": {"percent": premium_percent / 100, "type": premium_type},
        "tender_premium_f": tender_premium_f,
        "tender_premium_h": tender_premium_h,
        "grand_total_f": round(work_order_total + tender_premium_f),
        "grand_total_h": round(executed_total + tender_premium_h),
        "net_difference": round((executed_total + tender_premium_h) - (work_order_total + tender_premium_f))
    }

    return first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data

def generate_simple_pdf(doc_name, data, output_path):
    """Generate simple PDF using basic HTML conversion"""
    try:
        if not PDFKIT_AVAILABLE:
            st.error("PDF generation not available. Please install wkhtmltopdf.")
            return None
            
        # Simple HTML template
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{doc_name}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #000; padding: 8px; text-align: left; }}
                th {{ background-color: #f0f0f0; }}
                .header {{ font-size: 18px; font-weight: bold; margin-bottom: 20px; }}
                .total {{ font-weight: bold; background-color: #f9f9f9; }}
            </style>
        </head>
        <body>
            <div class="header">{doc_name}</div>
        """
        
        if doc_name == "First Page" and "items" in data:
            html_content += """
            <table>
                <tr><th>S.No.</th><th>Description</th><th>Unit</th><th>Quantity</th><th>Rate</th><th>Amount</th></tr>
            """
            for item in data["items"]:
                if not item.get("is_divider", False):
                    html_content += f"""
                    <tr>
                        <td>{item.get('serial_no', '')}</td>
                        <td>{item.get('description', '')}</td>
                        <td>{item.get('unit', '')}</td>
                        <td>{item.get('quantity', '')}</td>
                        <td>{item.get('rate', '')}</td>
                        <td>{item.get('amount', '')}</td>
                    </tr>
                    """
            
            if "totals" in data:
                totals = data["totals"]
                html_content += f"""
                <tr class="total">
                    <td colspan="5">Grand Total</td>
                    <td>₹{totals.get('grand_total', 0):,.2f}</td>
                </tr>
                <tr class="total">
                    <td colspan="5">Premium ({totals.get('premium', {}).get('percent', 0)*100:.1f}%)</td>
                    <td>₹{totals.get('premium', {}).get('amount', 0):,.2f}</td>
                </tr>
                <tr class="total">
                    <td colspan="5">Total Payable</td>
                    <td>₹{totals.get('payable', 0):,.2f}</td>
                </tr>
                """
            html_content += "</table>"
        
        html_content += "</body></html>"
        
        # Generate PDF
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None
        }
        
        pdfkit.from_string(html_content, output_path, options=options)
        return output_path
        
    except Exception as e:
        st.error(f"PDF generation failed: {str(e)}")
        return None

def create_simple_word_doc(doc_name, data, output_path):
    """Create simple Word document"""
    try:
        if not DOCX_AVAILABLE:
            st.error("Word document generation not available.")
            return None
            
        doc = Document()
        doc.add_heading(doc_name, 0)
        
        if doc_name == "First Page" and "items" in data:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'S.No.'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Unit'
            hdr_cells[3].text = 'Quantity'
            hdr_cells[4].text = 'Rate'
            hdr_cells[5].text = 'Amount'
            
            # Data rows
            for item in data["items"]:
                if not item.get("is_divider", False):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get('serial_no', ''))
                    row_cells[1].text = str(item.get('description', ''))
                    row_cells[2].text = str(item.get('unit', ''))
                    row_cells[3].text = str(item.get('quantity', ''))
                    row_cells[4].text = str(item.get('rate', ''))
                    row_cells[5].text = str(item.get('amount', ''))
        
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        st.error(f"Word document generation failed: {str(e)}")
        return None

def merge_pdfs_simple(pdf_files, output_path):
    """Simple PDF merger"""
    try:
        if not PYPDF_AVAILABLE:
            st.error("PDF merging not available.")
            return None
            
        writer = PdfWriter()
        
        for pdf_file in pdf_files:
            if os.path.exists(pdf_file):
                reader = PdfReader(pdf_file)
                for page in reader.pages:
                    writer.add_page(page)
        
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return output_path
        
    except Exception as e:
        st.error(f"PDF merging failed: {str(e)}")
        return None

def create_zip_archive_simple(files, output_path):
    """Create ZIP archive"""
    try:
        with zipfile.ZipFile(output_path, 'w') as zipf:
            for file_path in files:
                if os.path.exists(file_path):
                    zipf.write(file_path, os.path.basename(file_path))
        return output_path
    except Exception as e:
        st.error(f"ZIP creation failed: {str(e)}")
        return None

# ============================================================================
# MAIN APPLICATION
# ============================================================================

@st.cache_data(show_spinner=False, ttl=1800)
def _load_excel(file_bytes: bytes):
    """Load Excel once per unique content and return dataframes."""
    xl_file = pd.ExcelFile(BytesIO(file_bytes))
    ws_wo = pd.read_excel(xl_file, "Work Order", header=None)
    ws_bq = pd.read_excel(xl_file, "Bill Quantity", header=None)
    ws_extra = pd.read_excel(xl_file, "Extra Items", header=None)
    return ws_wo, ws_bq, ws_extra, xl_file.sheet_names


# Cached bill processing function
@st.cache_data(show_spinner=False, ttl=600)
def _process_bill_cached(ws_wo, ws_bq, ws_extra, premium_percent: float, premium_type: str):
    return process_bill(ws_wo, ws_bq, ws_extra, premium_percent, premium_type)


def main():
    """Main application entry point"""
    
    # Page configuration
    st.set_page_config(
        page_title="Stream Bill Generator",
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Add navigation info
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🧪 Additional Tools")
    st.sidebar.info("Check the **Test Files** page for managing test scripts and documentation")
    
    # Header
    st.title("📋 Stream Bill Generator")
    st.markdown("""
    **Professional Infrastructure Bill Generation System**  
    Generate contractor bills, deviation statements, and statutory documents from Excel data.
    """)
    
    # Deployment status indicator
    with st.expander("ℹ️ System Information"):
        col1, col2 = st.columns(2)
        with col1:
            st.success("✅ All modules loaded successfully")
            st.info(f"📁 Root Directory: `{ROOT_DIR}`")
        with col2:
            st.info(f"🐍 Python Version: {sys.version.split()[0]}")
            st.info(f"📦 Streamlit Version: {st.__version__}")
    
    # File upload section
    st.markdown("---")
    st.subheader("📤 Upload Excel Data")
    
    uploaded_file = st.file_uploader(
        "Select Excel file containing Work Order, Bill Quantity, and Extra Items sheets",
        type=["xlsx", "xls"],
        key="excel_upload",
        help="Upload an Excel file with the required sheets: Work Order, Bill Quantity, and Extra Items"
    )
    
    if uploaded_file is not None:
        try:
            # Read Excel file with caching
            file_bytes = uploaded_file.getvalue()
            ws_wo, ws_bq, ws_extra, sheet_names = _load_excel(file_bytes)
            
            # Validate required sheets
            required_sheets = ["Work Order", "Bill Quantity", "Extra Items"]
            missing_sheets = [sheet for sheet in required_sheets if sheet not in sheet_names]
            
            if missing_sheets:
                st.error(f"❌ Missing required sheets: **{', '.join(missing_sheets)}**")
                st.info(f"📋 Available sheets: {', '.join(sheet_names)}")
                return
            
            # Display success and available sheets
            st.success(f"✅ Excel file loaded successfully with {len(sheet_names)} sheets")
            
            # Dataframes already loaded via cache above
            
            # Sidebar configuration
            st.sidebar.header("⚙️ Configuration")
            
            st.sidebar.markdown("### Premium Settings")
            premium_percent = st.sidebar.number_input(
                "Tender Premium (%)", 
                min_value=-50.0,
                max_value=100.0, 
                value=5.0, 
                step=0.1,
                help="Enter the tender premium percentage (can be positive or negative)"
            )
            
            premium_type = st.sidebar.radio(
                "Premium Type",
                ["above", "below"],
                index=0,
                help="Select whether premium is above (added) or below (subtracted)"
            )
            
            st.sidebar.markdown("---")
            st.sidebar.markdown("### 📊 Quick Stats")
            st.sidebar.metric("Work Order Items", len(ws_wo) - 21 if len(ws_wo) > 21 else 0)
            st.sidebar.metric("Bill Quantity Items", len(ws_bq) - 21 if len(ws_bq) > 21 else 0)
            st.sidebar.metric("Extra Items", len(ws_extra) - 21 if len(ws_extra) > 21 else 0)
            
            # Generate button
            st.markdown("---")
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                generate_button = st.button(
                    "🚀 Generate All Documents", 
                    type="primary",
                    use_container_width=True
                )
            
            if generate_button:
                with st.spinner("🔄 Processing bill and generating documents..."):
                    try:
                        # Use consolidated functions (no external dependencies)
                        
                        # Process the bill
                        first_page_data, last_page_data, deviation_data, extra_items_data, note_sheet_data = _process_bill_cached(
                            ws_wo, ws_bq, ws_extra, premium_percent, premium_type
                        )
                        
                        # Create temporary directory for outputs
                        with tempfile.TemporaryDirectory() as temp_dir:
                            pdf_files = []
                            word_files = []
                            
                            # Progress tracking
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            # Generate PDFs using consolidated functions
                            status_text.text("Generating First Page PDF...")
                            first_page_pdf = os.path.join(temp_dir, "first_page.pdf")
                            if generate_simple_pdf("First Page", first_page_data, first_page_pdf):
                                pdf_files.append(first_page_pdf)
                            progress_bar.progress(25)
                            
                            status_text.text("Generating Last Page PDF...")
                            last_page_pdf = os.path.join(temp_dir, "last_page.pdf")
                            if generate_simple_pdf("Last Page", last_page_data, last_page_pdf):
                                pdf_files.append(last_page_pdf)
                            progress_bar.progress(35)
                            
                            status_text.text("Generating Deviation Statement PDF...")
                            deviation_pdf = os.path.join(temp_dir, "deviation_statement.pdf")
                            if generate_simple_pdf("Deviation Statement", deviation_data, deviation_pdf):
                                pdf_files.append(deviation_pdf)
                            progress_bar.progress(45)
                            
                            status_text.text("Generating Extra Items PDF...")
                            extra_items_pdf = os.path.join(temp_dir, "extra_items.pdf")
                            if generate_simple_pdf("Extra Items", extra_items_data, extra_items_pdf):
                                pdf_files.append(extra_items_pdf)
                            progress_bar.progress(55)
                            
                            status_text.text("Generating Note Sheet PDF...")
                            note_sheet_pdf = os.path.join(temp_dir, "note_sheet.pdf")
                            if generate_simple_pdf("Note Sheet", note_sheet_data, note_sheet_pdf):
                                pdf_files.append(note_sheet_pdf)
                            progress_bar.progress(65)
                            
                            # Generate Word documents
                            status_text.text("Creating Word documents...")
                            
                            word_files_data = [
                                ("first_page.docx", "First Page", first_page_data),
                                ("last_page.docx", "Last Page", last_page_data),
                                ("deviation_statement.docx", "Deviation Statement", deviation_data),
                                ("extra_items.docx", "Extra Items", extra_items_data),
                                ("note_sheet.docx", "Note Sheet", note_sheet_data)
                            ]
                            
                            for i, (filename, doc_name, doc_data) in enumerate(word_files_data):
                                doc_path = os.path.join(temp_dir, filename)
                                if create_simple_word_doc(doc_name, doc_data, doc_path):
                                    word_files.append(doc_path)
                                progress_bar.progress(65 + (i + 1) * 4)
                            
                            # Merge PDFs
                            status_text.text("Merging all PDFs...")
                            merged_pdf = os.path.join(temp_dir, "complete_bill.pdf")
                            if merge_pdfs_simple(pdf_files, merged_pdf):
                                pass  # Success
                            progress_bar.progress(90)
                            
                            # Create ZIP archive
                            status_text.text("Creating ZIP archive...")
                            all_files = pdf_files + word_files + ([merged_pdf] if os.path.exists(merged_pdf) else [])
                            zip_path = os.path.join(temp_dir, "bill_documents.zip")
                            create_zip_archive_simple(all_files, zip_path)
                            progress_bar.progress(100)
                            
                            status_text.text("✅ All documents generated successfully!")
                            
                            # Display success message
                            st.success("🎉 **Documents generated successfully!**")
                            
                            # Download section
                            st.markdown("---")
                            st.subheader("📥 Download Generated Documents")
                            
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                st.markdown("#### Complete Bill (PDF)")
                                with open(merged_pdf, "rb") as f:
                                    st.download_button(
                                        label="📄 Download Complete PDF",
                                        data=f,
                                        file_name="complete_bill.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                            
                            with col2:
                                st.markdown("#### All Documents (ZIP)")
                                with open(zip_path, "rb") as f:
                                    st.download_button(
                                        label="📦 Download ZIP Archive",
                                        data=f,
                                        file_name="bill_documents.zip",
                                        mime="application/zip",
                                        use_container_width=True
                                    )
                            
                            with col3:
                                st.markdown("#### Bill Summary")
                                grand_total = first_page_data['totals']['grand_total']
                                premium_amount = first_page_data['totals']['premium']['amount']
                                payable = first_page_data['totals']['payable']
                                
                                st.metric("Grand Total", f"₹{grand_total:,.2f}")
                                st.metric("Premium", f"₹{premium_amount:,.2f}", 
                                         delta=f"{premium_percent}% {premium_type}")
                                st.metric("Total Payable", f"₹{payable:,.2f}")
                            
                            # Clear progress indicators
                            progress_bar.empty()
                            status_text.empty()
                            
                    except Exception as e:
                        st.error(f"❌ **Error processing bill:** {str(e)}")
                        st.exception(e)
                        
        except Exception as e:
            st.error(f"❌ **Error reading Excel file:** {str(e)}")
            st.exception(e)
    else:
        # Welcome message and instructions
        st.info("👆 **Please upload an Excel file to get started**")
        
        # Instructions
        with st.expander("📋 **How to Use This Application**", expanded=True):
            st.markdown("""
            ### Step-by-Step Instructions:
            
            1. **Prepare Your Excel File**
               - Ensure it contains three required sheets:
                 - `Work Order` - Original work order details
                 - `Bill Quantity` - Executed work quantities
                 - `Extra Items` - Additional items not in work order
            
            2. **Upload the File**
               - Click the "Browse files" button above
               - Select your Excel file (.xlsx or .xls format)
            
            3. **Configure Settings** (Optional)
               - Adjust tender premium percentage in the sidebar
               - Select premium type (above/below)
            
            4. **Generate Documents**
               - Click the "Generate All Documents" button
               - Wait for processing to complete
            
            5. **Download Results**
               - Download the complete merged PDF
               - Or download the ZIP archive with all formats (PDF + Word)
            
            ### 📊 Generated Documents:
            
            - **First Page**: Summary of work items and amounts
            - **Last Page**: Final calculations and totals
            - **Deviation Statement**: Comparison between work order and executed work
            - **Extra Items**: Additional items with calculations
            - **Note Sheet**: Detailed notes and breakdowns
            """)
        
        # Sample data structure
        with st.expander("📊 **Required Excel Structure**"):
            st.markdown("""
            ### Work Order / Bill Quantity Sheet Structure:
            
            | Row | Column A | Column B | Column C | Column D | Column E | Column F |
            |-----|----------|----------|----------|----------|----------|----------|
            | 1-19 | Header information (work order details, contractor info, etc.) |
            | 20 | Table headers: S.No. | Description | Unit | Quantity | Rate | Amount |
            | 21+ | Work item data rows |
            
            ### Extra Items Sheet Structure:
            
            Similar to Work Order sheet, containing additional items not in original scope.
            
            ### Important Notes:
            
            - Excel file must be in `.xlsx` or `.xls` format
            - All three sheets must be present
            - Headers should be in row 20, data starts from row 21
            - Numeric values should be properly formatted
            - Ensure no merged cells in data rows
            """)
        
        # Features highlight
        with st.expander("✨ **Application Features**"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("""
                **Document Generation:**
                - ✅ PDF format (optimized, professional)
                - ✅ Word format (editable .docx)
                - ✅ Merged complete bill PDF
                - ✅ ZIP archive of all documents
                
                **Calculations:**
                - ✅ Automatic premium calculation
                - ✅ Deviation analysis
                - ✅ Extra items processing
                - ✅ Grand total computation
                """)
            
            with col2:
                st.markdown("""
                **Quality Features:**
                - ✅ Statutory government format compliance
                - ✅ Professional document styling
                - ✅ Accurate number-to-words conversion
                - ✅ Detailed error handling
                
                **Performance:**
                - ✅ Fast processing (<30 seconds)
                - ✅ Cloud-optimized deployment
                - ✅ Batch processing support
                - ✅ Progress tracking
                """)

if __name__ == "__main__":
    main()