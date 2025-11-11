"""
Word Document Generator
Generates editable .docx files from bill data for easy last-minute edits
"""
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    
    tcPr.append(tcBorders)

def generate_first_page_docx(data, output_path):
    """Generate First Page as Word document"""
    doc = Document()
    
    # Set margins (10mm all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.page_width = Mm(210)  # A4
        section.page_height = Mm(297)
    
    # Title
    title = doc.add_heading('CONTRACTOR BILL', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Header information
    for row in data.get('header', []):
        if row and any(str(item).strip() for item in row if item):
            p = doc.add_paragraph()
            p.style = 'Normal'
            for item in row:
                if item and str(item).strip():
                    p.add_run(str(item) + ' ')
    
    # Main table
    items = data.get('items', [])
    if items:
        # Create table with 9 columns
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Unit', 'Qty Since Last', 'Qty Upto Date', 'S.No.', 
                   'Description', 'Rate', 'Amount Upto Date', 'Amount Since Previous', 'Remarks']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
        
        # Data rows
        for item in items:
            if item.get('is_divider'):
                row = table.add_row()
                row.cells[0].merge(row.cells[8])
                row.cells[0].text = item.get('description', '')
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            else:
                row = table.add_row()
                row.cells[0].text = str(item.get('unit', ''))
                row.cells[1].text = str(item.get('quantity_since_last', ''))
                row.cells[2].text = str(item.get('quantity_upto_date', ''))
                row.cells[3].text = str(item.get('serial_no', ''))
                row.cells[4].text = str(item.get('description', ''))
                row.cells[5].text = str(item.get('rate', ''))
                row.cells[6].text = str(item.get('amount', ''))
                row.cells[7].text = str(item.get('amount_previous', ''))
                row.cells[8].text = str(item.get('remark', ''))
        
        # Totals
        totals = data.get('totals', {})
        
        # Grand Total
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        row.cells[0].text = 'Grand Total Rs.'
        row.cells[6].text = str(totals.get('grand_total', ''))
        row.cells[7].text = str(totals.get('grand_total', ''))
        
        # Premium
        premium = totals.get('premium', {})
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        row.cells[0].text = f"Tender Premium @ {premium.get('percent', 0)*100:.2f}%"
        row.cells[6].text = str(premium.get('amount', ''))
        row.cells[7].text = str(premium.get('amount', ''))
        
        # Extra Items Sum
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        extra_sum = totals.get('extra_items_sum', 0)
        row.cells[0].text = str(extra_sum if extra_sum > 0 else 'NIL')
        row.cells[4].text = 'Sum of Extra Items (including Tender Premium) Rs.'
        
        # Payable Amount
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        row.cells[0].text = 'Payable Amount Rs.'
        row.cells[6].text = str(totals.get('payable', ''))
        row.cells[7].text = str(totals.get('payable', ''))
        
        # Less Amount Paid
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        row.cells[0].text = 'Less Amount Paid vide Last Bill Rs.'
        last_bill = totals.get('last_bill_amount', 0)
        row.cells[6].text = f"{last_bill:.2f}" if last_bill else "0.00"
        row.cells[7].text = f"{last_bill:.2f}" if last_bill else "0.00"
        
        # Net Payable
        row = table.add_row()
        row.cells[0].merge(row.cells[3])
        row.cells[0].text = 'Net Payable Amount Rs.'
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        net_payable = totals.get('net_payable', totals.get('payable', 0))
        row.cells[6].text = f"{net_payable:.2f}"
        row.cells[7].text = f"{net_payable:.2f}"
        row.cells[6].paragraphs[0].runs[0].font.bold = True
        row.cells[7].paragraphs[0].runs[0].font.bold = True
    
    # Save
    doc.save(output_path)
    return output_path

def generate_deviation_statement_docx(data, output_path):
    """Generate Deviation Statement as Word document"""
    doc = Document()
    
    # Set margins and landscape
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.page_width = Mm(297)  # A4 Landscape
        section.page_height = Mm(210)
    
    # Title
    title = doc.add_heading('Deviation Statement', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Main table
    items = data.get('items', [])
    if items:
        table = doc.add_table(rows=1, cols=13)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['ITEM No.', 'Description', 'Unit', 'Qty WO', 'Rate', 'Amt WO', 
                   'Qty Executed', 'Amt Executed', 'Excess Qty', 'Excess Amt', 
                   'Saving Qty', 'Saving Amt', 'Remarks']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
        
        # Data rows
        for item in items:
            if item.get('is_divider'):
                row = table.add_row()
                row.cells[0].merge(row.cells[12])
                row.cells[0].text = item.get('description', '')
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row = table.add_row()
                row.cells[0].text = str(item.get('serial_no', ''))
                row.cells[1].text = str(item.get('description', ''))
                row.cells[2].text = str(item.get('unit', ''))
                row.cells[3].text = str(item.get('qty_wo', ''))
                row.cells[4].text = str(item.get('rate', ''))
                row.cells[5].text = str(item.get('amt_wo', ''))
                row.cells[6].text = str(item.get('qty_bill', ''))
                row.cells[7].text = str(item.get('amt_bill', ''))
                row.cells[8].text = str(item.get('excess_qty', ''))
                row.cells[9].text = str(item.get('excess_amt', ''))
                row.cells[10].text = str(item.get('saving_qty', ''))
                row.cells[11].text = str(item.get('saving_amt', ''))
                row.cells[12].text = str(item.get('remark', ''))
        
        # Summary rows
        summary = data.get('summary', {})
        
        # Grand Total
        row = table.add_row()
        row.cells[0].merge(row.cells[4])
        row.cells[0].text = 'Grand Total Rs.'
        row.cells[5].text = str(summary.get('work_order_total', ''))
        row.cells[7].text = str(summary.get('executed_total', ''))
        row.cells[9].text = str(summary.get('overall_excess', ''))
        row.cells[11].text = str(summary.get('overall_saving', ''))
        
        # Premium
        premium = summary.get('premium', {})
        row = table.add_row()
        row.cells[0].merge(row.cells[4])
        row.cells[0].text = f"Add Tender Premium ({premium.get('percent', 0)*100:.2f}%)"
        row.cells[5].text = str(summary.get('tender_premium_f', ''))
        row.cells[7].text = str(summary.get('tender_premium_h', ''))
        row.cells[9].text = str(summary.get('tender_premium_j', ''))
        row.cells[11].text = str(summary.get('tender_premium_l', ''))
        
        # Grand Total with Premium
        row = table.add_row()
        row.cells[0].merge(row.cells[4])
        row.cells[0].text = 'Grand Total including Tender Premium Rs.'
        row.cells[5].text = str(summary.get('grand_total_f', ''))
        row.cells[7].text = str(summary.get('grand_total_h', ''))
        row.cells[9].text = str(summary.get('grand_total_j', ''))
        row.cells[11].text = str(summary.get('grand_total_l', ''))
        
        # Net Difference
        row = table.add_row()
        row.cells[0].merge(row.cells[6])
        is_saving = summary.get('is_saving', False)
        row.cells[0].text = 'Overall Saving With Respect to the Work Order Amount Rs.' if is_saving else 'Overall Excess With Respect to the Work Order Amount Rs.'
        row.cells[7].text = str(summary.get('net_difference', ''))
        
        # Percentage Deviation
        row = table.add_row()
        row.cells[0].merge(row.cells[6])
        row.cells[0].text = 'Percentage of Deviation %'
        row.cells[7].text = f"{summary.get('percentage_deviation', 0):.2f}%"
    
    # Save
    doc.save(output_path)
    return output_path

def generate_extra_items_docx(data, output_path):
    """Generate Extra Items as Word document"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
    
    # Title
    title = doc.add_heading('Extra Items', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Main table
    items = data.get('items', [])
    if items:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['S.No.', 'Remarks', 'Description', 'Quantity', 'Unit', 'Rate', 'Amount']
        
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
        
        # Data rows
        for item in items:
            row = table.add_row()
            row.cells[0].text = str(item.get('serial_no', ''))
            row.cells[1].text = str(item.get('remark', ''))
            row.cells[2].text = str(item.get('description', ''))
            row.cells[3].text = str(item.get('quantity', ''))
            row.cells[4].text = str(item.get('unit', ''))
            row.cells[5].text = str(item.get('rate', ''))
            row.cells[6].text = str(item.get('amount', ''))
    
    # Save
    doc.save(output_path)
    return output_path
